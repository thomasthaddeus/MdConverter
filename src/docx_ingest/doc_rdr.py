"""
_summary_

_extended_summary_
"""

import os
from docx import Document
import zipfile
import xml.etree.ElementTree as ET
import shutil


class DocumentReader:
    """
    A class used to read a Word document.

    ...

    Attributes
    ----------
    file_path : str
        the path to the Word document
    document : Document
        the Word document

    Methods
    -------
    read_document():
        Reads the Word document.
    """

    def __init__(self, file_path):
        self.file_path = file_path
        self.document = None

    def read_document(self):
        """
        Reads the Word document.

        Raises
        ------
        Exception
            If the file does not exist, is not a valid Word document, or is empty.
        """
        if not os.path.exists(self.file_path):
            raise Exception(f"File does not exist: {self.file_path}")

        try:
            self.document = Document(self.file_path)
        except Exception as exc:
            raise Exception(f"Error reading document: {exc}")

        if len(self.document.paragraphs) == 0 and len(self.document.tables) == 0:
            raise Exception("Document is empty")

        contains_elements_to_parse = False
        for paragraph in self.document.paragraphs:
            if paragraph.style.name in ["Heading 1", "Normal", "List Paragraph"]:
                contains_elements_to_parse = True
                break

        if not contains_elements_to_parse and len(self.document.tables) == 0:
            raise Exception("Document does not contain any elements to parse")

    def parse_raw_xml(self):
        # Save a copy of the document
        self.document.save('copy.docx')

        # Unzip the document
        with zipfile.ZipFile('copy.docx', 'r') as zip_ref:
            zip_ref.extractall('copy')

        # Parse the document XML
        tree = ET.parse('copy/word/document.xml')
        root = tree.getroot()

        # Define the WordprocessingML namespace
        namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        # Find the list items
        list_items = root.findall('.//w:p[w:pPr/w:numPr]', namespaces)
        for list_item in list_items:
            print(ET.tostring(list_item, encoding='unicode'))

    def extract_images(self):
        # Create the img directory if it doesn't exist
        os.makedirs('img', exist_ok=True)

        # Copy the image files to the img directory
        for filename in os.listdir('copy/word/media'):
            shutil.copy(os.path.join('copy/word/media', filename), 'img')
