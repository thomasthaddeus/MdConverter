"""
_summary_

_extended_summary_
"""


from docx import Document
from docx.shared import Pt
import os
import shutil

class DocumentParser:
    """
    A class used to parse different elements of a Word document.

    ...

    Attributes
    ----------
    document : Document
        a Word document to parse
    headings : list
        a list of headings in the document
    paragraphs : list
        a list of paragraphs in the document
    lists : list
        a list of lists in the document
    tables : list
        a list of tables in the document
    images : list
        a list of images in the document

    Methods
    -------
    parse_headings():
        Parses the headings in the document.
    parse_paragraphs():
        Parses the paragraphs in the document.
    parse_lists():
        Parses the lists in the document.
    parse_tables():
        Parses the tables in the document.
    parse_images():
        Parses the images in the document.
    """

    def __init__(self, document):
        self.document = document
        self.headings = []
        self.paragraphs = []
        self.lists = []
        self.tables = []
        self.images = []

    def parse_headings(self):
        """
        Parses the headings in the document.

        Raises
        ------
        Exception
            If the document is None.
        """
        if self.document is None:
            raise Exception("Document is None")

        for paragraph in self.document.paragraphs:
            if paragraph.style.name == "Heading 1":
                self.headings.append(paragraph.text)

    def parse_paragraphs(self):
        """
        Parses the paragraphs in the document.

        Raises
        ------
        Exception
            If the document is None.
        """
        if self.document is None:
            raise Exception("Document is None")

        for paragraph in self.document.paragraphs:
            if paragraph.style.name == "Normal":
                self.paragraphs.append(paragraph.text)

    def parse_lists(self):
        """
        Parses the lists in the document.

        Raises
        ------
        Exception
            If the document is None.
        """
        if self.document is None:
            raise Exception("Document is None")

        for paragraph in self.document.paragraphs:
            if paragraph.style.name == "List Paragraph":
                self.lists.append(paragraph.text)

    def parse_tables(self):
        """
        Parses the tables in the document.

        Raises
        ------
        Exception
            If the document is None.
        """
        if self.document is None:
            raise Exception("Document is None")

        for table in self.document.tables:
            self.tables.append(table)

    def parse_images(self):
        """
        Parses the images in the document.

        Raises
        ------
        Exception
            If the document is None.
        """
        if self.document is None:
            raise Exception("Document is None")

        for rel in self.document.part.rels.values():
            if "image" in rel.reltype:
                image_data = rel.blob
                image_path = os.path.join("/tmp", rel.relid)
                with open(image_path, 'wb') as f:
                    f.write(image_data)

                self.images.append(image_path)
